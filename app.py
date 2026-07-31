import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from groq import Groq

resume_file_path = ""


def extract_job_title(soup):
    # Just use the page's <title> tag as-is
    if soup.title and soup.title.string:
        return soup.title.string.strip()
    return None


def sanitize_filename(name):
    # Strip characters that aren't safe in Windows/macOS/Linux filenames
    invalid_chars = '\\/:*?"<>|'
    cleaned = "".join(c for c in name if c not in invalid_chars)
    cleaned = cleaned.strip().replace(" ", "_")
    return cleaned or "Job"


def scraper(link):
    # Fetch and extract job description text + job title from job listing link
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }

    try:
        response = requests.get(link, headers=headers)
        response.raise_for_status()
    except requests.RequestException as e:
        print(f"Failed to fetch page: {e}")
        return None, None

    soup = BeautifulSoup(response.content, "html.parser")
    posting = soup.find('div', class_='show-more-less-html__markup')
    job_title = extract_job_title(soup)

    if posting:
        text_content = posting.get_text(separator=' ', strip=True)
        return text_content, job_title
    else:
        print("Could not find the target job description container on this page.")
        return None, job_title


def process_single_bullet(client, bullet_text, job_description):
    # Send an individual bullet point to Groq for targeted keyphrase optimization
    prompt = f"""
    JOB DESCRIPTION:
    {job_description}

    ORIGINAL RESUME BULLET:
    {bullet_text}

    INSTRUCTIONS:
    Rewrite this single bullet point to better align with the job description keywords while maintaining factual accuracy to the original text. If the original description has nothing to do with the new job posting, focus on explaining the soft people skills learned. 
    Output ONLY the rewritten bullet point text—no introductory chat, quotes, or markdown formatting. Output should only be 230 characters at most. 
    """

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )

    return response.choices[0].message.content.strip()


def generator(job_description, job_title=None):
    global resume_file_path

    if not os.path.isfile(resume_file_path):
        print("Resume path not found. Make sure get_save() has run first.")
        return None

    doc = Document(resume_file_path)

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        api_key = input("Please enter your Groq API Key :: ").strip()

    client = Groq(api_key=api_key)

    print("\nScanning document for bullet points...")
    bullets_found = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = str(para.style.name).lower() if para.style else ""

        # Check for real Word list/numbering formatting (w:numPr in the XML),
        # rather than guessing from indentation, which also matches
        # non-bullet paragraphs like titles or addresses.
        has_numbering = para._p.find(
            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr'
        ) is not None

        # Bullet detection logic
        is_bullet = (
            "bullet" in style_name
            or "list" in style_name
            or has_numbering
            or text.startswith("•")
            or text.startswith("-")
            or text.startswith("*")
            or text.startswith("▪")
        )

        # Ignore short text like section titles or dates
        if is_bullet and len(text) > 15:
            bullets_found += 1
            print(f"\n[Bullet #{bullets_found}] Original: {text}")

            tailored_bullet = process_single_bullet(client, text, job_description)
            print(f"[Bullet #{bullets_found}] Tailored: {tailored_bullet}")

            # Overwrite text directly in paragraph
            if para.runs:
                first_run = para.runs[0]
                is_bold = first_run.bold
                font_name = first_run.font.name
                font_size = first_run.font.size

                para.text = tailored_bullet

                if para.runs:
                    para.runs[0].bold = is_bold
                    if font_name:
                        para.runs[0].font.name = font_name
                    if font_size:
                        para.runs[0].font.size = font_size
            else:
                para.text = tailored_bullet

    if bullets_found == 0:
        print("\nNo bullets detected! Your template might use tables or unindented paragraphs.")
        print("Let's print out what python-docx sees in your document to fix the filter:")
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                print(f" Line {i}: '{para.text.strip()}' | Style: '{para.style.name}'")

    original_name = os.path.splitext(os.path.basename(resume_file_path))[0]

    if job_title:
        output_filename = f"{original_name}_{sanitize_filename(job_title)}.docx"
    else:
        output_filename = f"{original_name}_Tailored.docx"

    doc.save(output_filename)
    print(f"\nSuccessfully generated updated resume at: {output_filename}")
    return output_filename


def get_save():
    # Load or prompt for base resume path and store in config/global
    global resume_file_path
    config = "pathResume.txt"

    if os.path.isfile(config):
        with open(config, "r") as file:
            resume_file_path = file.read().strip()
            print(f"Found resume path:: {resume_file_path}")
            return resume_file_path
    while True:
        user_path = input("Please input the path to your base resume :: ")
        if os.path.isfile(user_path) and user_path.endswith(".docx"):
            with open(config, "w") as file:
                file.write(user_path)
                resume_file_path = user_path
                return user_path
        else:
            print("Invalid path or file is not a docx file!")


def main():
    base_path = get_save()
    doc = Document(base_path)
    print("Resume successfully loaded")

    link = input("Please input job link :: ")
    job_text, job_title = scraper(link)

    if job_text:
        generator(job_text, job_title)
    else:
        print("Could not extract job text. Skipping generation step.")


if __name__ == "__main__":
    main()